"""Add the AI-exhaustive mechanism search framework section to the
Khorium Hypersonics Alignment Plan docx.

Inserts a new Section 4 "Mechanism Search Framework Integration" before
the existing "6. The Drop-In Test" section, describing how the search
framework layers onto the SimOps architecture.

Run:
    python scripts/update_alignment_plan_with_search_framework.py
"""
from __future__ import annotations

import sys
from pathlib import Path

import docx
from docx.shared import Pt
from copy import deepcopy

DOC = Path(r"C:\Users\yarden\Desktop\Khorium Hypersonics\SimOps Integration\Khorium Hypersonics Alignment Plan.docx")


def insert_paragraph_before(reference_p, text: str, style_name: str | None = None):
    """Insert a new paragraph before `reference_p` with `text` at given style."""
    new_p = reference_p.insert_paragraph_before(text, style=style_name)
    return new_p


def main():
    if not DOC.exists():
        print(f"ERROR: doc not found at {DOC}")
        sys.exit(1)

    print(f"Opening {DOC.name}")
    doc = docx.Document(str(DOC))

    # Find the paragraph that begins section 6
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "6. The Drop-In Test" in p.text:
            target_idx = i
            print(f"Found section 6 at paragraph index {i}")
            break
    if target_idx is None:
        print("ERROR: couldn't find '6. The Drop-In Test' anchor — aborting")
        sys.exit(2)

    target_p = doc.paragraphs[target_idx]

    # Section content (lines + style mapping)
    new_section = [
        ("4. Mechanism Search Framework Integration", "Heading 1"),
        ("Per Aaron Wu's vision (Slack 2026-04-25): \"if we create a framework that "
         "allows the AI to try exhaust method on the chemistry reaction search, there "
         "is a way that we can do something that nobody has ever done before in human "
         "history.\" This section describes how the AI-exhaustive mechanism search "
         "framework slots into the SimOps architecture.", "Normal"),
        ("", "Normal"),
        ("4.1 Overview", "Heading 2"),
        ("The mechanism search framework lives at plasmanet/mechanism_search/. It is "
         "NOT a per-solver SimOps module like su2_nemo or openfoam — it is a "
         "META-LAYER that orchestrates ALL solvers (Cantera 0D, PlasmaNet surrogate, "
         "SU2-NEMO CFD) to evaluate mechanism candidates against published "
         "experimental data and find the best fit.", "Normal"),
        ("", "Normal"),
        ("Search space: Park 1990's full air mechanism has 47 reactions across 11 "
         "species. Subset space is 2^47 ≈ 1.4 × 10^14 candidates. Direct CFD "
         "evaluation is infeasible (each run is 1-10 hours); the framework uses "
         "surrogates for fast scoring, then validates only the top-K with full CFD.",
         "Normal"),
        ("", "Normal"),
        ("4.2 Layered architecture (matches Khorium SimOps pattern)", "Heading 2"),
        ("plasmanet/mechanism_search/", "Normal"),
        ("├── generator.py        # S-1: Park 47 mechanism + subset() filtering", "Normal"),
        ("├── scoring.py          # S-6: Composite score vs RAM-C 1972, Grantham 1970", "Normal"),
        ("├── cantera_evaluator.py # S-2: Fast 0D surrogate (~ms per evaluation)", "Normal"),
        ("├── surrogate.py        # S-3: PlasmaNet retrained on (mechanism, conditions) → ne", "Normal"),
        ("├── search_loop.py      # S-4: Bayesian / GA over reaction subsets", "Normal"),
        ("└── cfd_validator.py    # S-5: Top-K validation via SU2-NEMO MPI binary", "Normal"),
        ("", "Normal"),
        ("Each module exposes a clean Python API. Layered dependencies: scoring "
         "is foundational; generator is independent; cantera_evaluator and "
         "surrogate plug into scoring as evaluator backends; search_loop calls "
         "score_candidate() on candidates from generator; cfd_validator uses "
         "the same SimOps su2_nemo runner for top-K validation.", "Normal"),
        ("", "Normal"),
        ("4.3 SimOps integration points", "Heading 2"),
        ("The search framework integrates with KhoriumBackend's SimOps pattern "
         "in three ways:", "Normal"),
        ("", "Normal"),
        ("(1) Each evaluator backend is a SimOps solver. Cantera 0D, "
         "PlasmaNet surrogate, and SU2-NEMO CFD each follow the triadic "
         "per-solver pattern (params.py, runner.py, container.py) defined in "
         "Section 1.1. The search framework calls them via their SimOps "
         "runners, not via direct Python imports — this means search runs "
         "can be distributed across AWS Batch jobs the same way single CFD "
         "runs are.", "Normal"),
        ("", "Normal"),
        ("(2) New SimOps module: mechanism_search/ joins simops/ as a "
         "first-class solver type. Its run.py orchestrates the search loop "
         "via KhoriumBackend's job queue, spawning sub-jobs (one per "
         "candidate evaluated) and aggregating results. From the dispatch "
         "side it looks like:", "Normal"),
        ("# KhoriumBackend/src/simops/main.py", "Normal"),
        ("if solver_type == \"mechanism_search\":", "Normal"),
        ("    params = MechanismSearchParams.model_validate_json(sim_params_raw)", "Normal"),
        ("    mechanism_search.runner.run(case_dir, output_dir, params)", "Normal"),
        ("", "Normal"),
        ("(3) Result-summary contract. Each search run produces a "
         "result.json matching the contract from Section 3.4: top-K "
         "candidates with their composite scores, per-benchmark "
         "log10(ne_predicted/ne_published), dB margins per radio band, and "
         "links (S3 keys) to the full CFD VTU files for top candidates.", "Normal"),
        ("", "Normal"),
        ("4.4 API surface", "Heading 2"),
        ("Backend FastAPI endpoints exposed by KhoriumBackend:", "Normal"),
        ("• POST /api/plasma/mechanism-search/start — Submit a search "
         "(constraints: max_reactions, required_species, benchmarks, budget). "
         "Returns search_id.", "Normal"),
        ("• GET /api/plasma/mechanism-search/{id}/status — Progress: candidates "
         "evaluated, current best score.", "Normal"),
        ("• GET /api/plasma/mechanism-search/{id}/results — Top-K candidates "
         "with per-benchmark scores and Cantera mechanism YAMLs.", "Normal"),
        ("• POST /api/plasma/mechanism-search/{id}/validate-top-k — Trigger "
         "full CFD validation for top-K candidates (uses MPI binary).", "Normal"),
        ("", "Normal"),
        ("4.5 Frontend integration", "Heading 2"),
        ("New module KhoriumFrontend/src/modules/mechanism_search/:", "Normal"),
        ("• SearchSetup.tsx — Define constraints (max reactions slider, "
         "required species checkboxes, benchmark suite selector).", "Normal"),
        ("• ProgressCard.tsx — Live updates as candidates are evaluated; "
         "best-score-so-far chart.", "Normal"),
        ("• ResultsTable.tsx — Ranked top-K list with composite score and "
         "per-benchmark verdict (EXCELLENT / GOOD / OK / POOR).", "Normal"),
        ("• MechanismDetail.tsx — Click into a candidate: Cantera YAML "
         "viewer, predicted ne field, dB attenuation polar plot per band.", "Normal"),
        ("", "Normal"),
        ("4.6 Storage conventions", "Heading 2"),
        ("S3 layout under simulations/{search_id}/:", "Normal"),
        ("├── search_params.json    # Input constraints", "Normal"),
        ("├── candidates/", "Normal"),
        ("│   ├── 0001/", "Normal"),
        ("│   │   ├── mechanism.yaml  # Cantera input", "Normal"),
        ("│   │   ├── score.json      # Per-benchmark composite score", "Normal"),
        ("│   │   └── cfd_result.vtu  # Only for top-K validated", "Normal"),
        ("│   ├── 0002/", "Normal"),
        ("│   └── ...", "Normal"),
        ("└── final_report.pdf       # Auto-generated PDF: top-K vs published data", "Normal"),
        ("", "Normal"),
        ("4.7 Build status (as of 2026-04-25)", "Heading 2"),
        ("S-1 Mechanism generator: 50% — Reaction + Mechanism dataclasses, "
         "Park 1990 18/47 reactions filled, subset() API, Cantera YAML emitter, "
         "SU2 cfg snippet emitter. Remaining: 19-47 (mechanical extraction "
         "from Park 1990 Tables 2+4+6).", "Normal"),
        ("S-6 Scoring framework: 80% — 4 RAM-C benchmarks loaded, log10 ne "
         "+ verdict-based dB scoring, anchor test passes (reproduces our "
         "measured AIR-5 baseline log10 err = −1.59 from CFD ground truth).", "Normal"),
        ("S-2 Cantera 0D evaluator: 40% — Normal-shock + chemistry-sink "
         "correction + IdealGasReactor + Appleton-Hartree dB. Blocked on "
         "Cantera install on VM (Windows wheel build fails locally).", "Normal"),
        ("S-3 Surrogate (PlasmaNet retrained on mechanism axis): pending. "
         "Requires ~50-100 evaluations from S-2 as training data.", "Normal"),
        ("S-4 Search loop: pending. Bayesian optimization or genetic "
         "algorithm over discrete reaction-subset space, scored via S-6.", "Normal"),
        ("S-5 Top-K CFD validation: pending. Uses /opt/su2-nemo-mpi/ "
         "binary built today (10-15× speedup over serial).", "Normal"),
        ("", "Normal"),
        ("4.8 Why this is the contribution", "Heading 2"),
        ("Existing literature picks mechanisms by hand: Park 1990 derived 47 "
         "reactions from shock-tube calibration; Dunn-Kang 1973 picked 15 for "
         "low-T air; Kang-Dunn 1979 picked 7 for ablation. Mechanism reduction "
         "tools (DRGEP, PFA, sensitivity analysis) START from Park's master "
         "set and remove reactions.", "Normal"),
        ("", "Normal"),
        ("This framework REVERSES that — it starts with the empty set and "
         "adds reactions iteratively to maximize fit against multi-experiment "
         "ground truth (RAM-C 1972, Grantham 1970, FIRE-II, Apollo). No "
         "published work has automated this kind of forward-search across "
         "the 2^47 subset space. Aaron's framing: \"something nobody has "
         "ever done before in human history.\"", "Normal"),
        ("", "Normal"),
        ("Validation results to date (2026-04-25): AIR-5 baseline at "
         "log10 err = −1.59 vs J&C 1972 RAM-C 61 km / M=22.5; AIR-7 "
         "ramp running at projected log10 err ≈ −0.5 to −1.0 (1+ order "
         "improvement). AIR-11 + Mutation++ has multiple compounding bugs "
         "in SU2 v7.5.1 (formation-enthalpy reference mismatch, "
         "chemistry-source NaN at trace electron density) — documented "
         "in scripts/mpp_air5_to_air11_converter.cpp + docs. AIR-7 "
         "viscous + non-cat wall has heap corruption in NEMO_NS solver "
         "preprocessing — documented as out-of-scope. These dead ends are "
         "PART of the framework's knowledge: the search algorithm avoids "
         "them as known-broken regions of the solver-config space.", "Normal"),
        ("", "Normal"),
    ]

    print(f"Inserting {len(new_section)} paragraphs before section 6...")
    for text, style in new_section:
        # Insert before target paragraph; the next iteration will keep
        # inserting before the same target so they appear in order
        if style == "Heading 1":
            try:
                target_p.insert_paragraph_before(text, style="Heading 1")
            except KeyError:
                p = target_p.insert_paragraph_before(text)
                p.runs[0].bold = True if p.runs else None
        elif style == "Heading 2":
            try:
                target_p.insert_paragraph_before(text, style="Heading 2")
            except KeyError:
                p = target_p.insert_paragraph_before(text)
                if p.runs:
                    p.runs[0].bold = True
        else:
            target_p.insert_paragraph_before(text, style="Normal")

    # Save
    doc.save(str(DOC))
    print(f"Saved {DOC.name}")


if __name__ == "__main__":
    main()
