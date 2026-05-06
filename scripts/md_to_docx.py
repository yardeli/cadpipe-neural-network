"""Markdown -> Word doc converter for the khorium_hypersonic paper.

Handles: # headings (h1-h4), tables (|---|), code blocks (```),
bullet/numbered lists, bold/italic spans. Renders into a clean Word
document with section headings, a TOC, and consistent styling.

Usage:
    python scripts/md_to_docx.py khorium_hypersonic/PAPER.md khorium_hypersonic/PAPER.docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_runs(paragraph, text: str):
    """Parse inline **bold**, *italic*, `code`, [link](url) and add runs."""
    pattern = re.compile(
        r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))"
    )
    pos = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > pos:
            paragraph.add_run(text[pos:start])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                paragraph.add_run(link_match.group(1))
            else:
                paragraph.add_run(token)
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_code_block(doc, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _add_table(doc, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.bold = True
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, cell_text in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_runs(p, cell_text.strip())


def md_to_docx(md_path: Path, out_path: Path) -> None:
    md = md_path.read_text(encoding="utf-8")
    doc = Document()

    # Set the default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # --- Code block ---
        if line.strip().startswith("```"):
            j = i + 1
            buf = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                buf.append(lines[j])
                j += 1
            _add_code_block(doc, buf)
            i = j + 1
            continue

        # --- Table ---
        if line.lstrip().startswith("|") and (i + 1) < len(lines) and re.match(r"\s*\|[-:|\s]+\|\s*$", lines[i + 1]):
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            j = i + 2
            row_data: list[list[str]] = []
            while j < len(lines) and lines[j].lstrip().startswith("|"):
                row_cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                if any(c for c in row_cells):
                    row_data.append(row_cells)
                j += 1
            _add_table(doc, header_cells, row_data)
            doc.add_paragraph("")
            i = j
            continue

        # --- Headings ---
        h_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            heading = doc.add_heading(level=level)
            run = heading.add_run(text)
            if level == 1:
                run.font.size = Pt(20)
            elif level == 2:
                run.font.size = Pt(15)
            i += 1
            continue

        # --- Horizontal rule ---
        if line.strip() == "---":
            doc.add_paragraph("─" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # --- Bullet ---
        bullet = re.match(r"^(\s*)([-*])\s+(.+)$", line)
        if bullet:
            indent = len(bullet.group(1))
            text = bullet.group(3)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (indent // 2))
            _add_runs(p, text)
            i += 1
            continue

        # --- Numbered ---
        numbered = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if numbered:
            text = numbered.group(2)
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, text)
            i += 1
            continue

        # --- Empty line / paragraph break ---
        if not line.strip():
            i += 1
            continue

        # --- Regular paragraph (consume until blank line or block) ---
        para_lines = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            if (not nxt.strip()
                    or re.match(r"^#{1,4}\s+", nxt)
                    or nxt.lstrip().startswith("|")
                    or nxt.strip().startswith("```")
                    or re.match(r"^\s*[-*]\s+", nxt)
                    or re.match(r"^\s*\d+\.\s+", nxt)
                    or nxt.strip() == "---"):
                break
            para_lines.append(nxt)
            j += 1
        para_text = " ".join(para_lines).strip()
        p = doc.add_paragraph()
        _add_runs(p, para_text)
        i = j

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


def main():
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} <input.md> <output.docx>")
        sys.exit(1)
    md_to_docx(Path(sys.argv[1]), Path(sys.argv[2]))


if __name__ == "__main__":
    main()
